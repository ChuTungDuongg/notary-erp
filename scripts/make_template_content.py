from pathlib import Path
from docx import Document

TEMPLATE_PATH = Path("templates/contract_transfer.docx")

def main():
    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("HỢP ĐỒNG CHUYỂN NHƯỢNG", level=1)

    doc.add_paragraph("BÊN BÁN: {{seller_name}}")
    doc.add_paragraph("CCCD: {{seller_cccd}}")
    doc.add_paragraph("ĐỊA CHỈ: {{seller_address}}")

    doc.add_paragraph("")

    doc.add_paragraph("BÊN MUA: {{buyer_name}}")
    doc.add_paragraph("CCCD: {{buyer_cccd}}")
    doc.add_paragraph("ĐỊA CHỈ: {{buyer_address}}")

    doc.add_paragraph("")
    doc.add_paragraph("TÀI SẢN:")
    doc.add_paragraph("- Địa chỉ: {{property_address}}")
    doc.add_paragraph("- Tờ bản đồ: {{property_map_sheet_no}}")
    doc.add_paragraph("- Thửa: {{property_parcel_no}}")
    doc.add_paragraph("- Diện tích: {{property_area_m2}} m2")
    doc.add_paragraph("- Số GCN: {{property_certificate_no}}")

    doc.add_paragraph("")
    doc.add_paragraph("GIÁ CHUYỂN NHƯỢNG: {{transfer_price}}")
    doc.add_paragraph("NGÀY KÝ: {{signing_date}}")

    doc.save(str(TEMPLATE_PATH))
    print(f"OK: wrote {TEMPLATE_PATH}")

if __name__ == "__main__":
    main()
