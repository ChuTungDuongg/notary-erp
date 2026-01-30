from __future__ import annotations

from pathlib import Path
from typing import Dict

from docx import Document as DocxDocument
from sqlalchemy import select, func
from sqlalchemy.orm import Session

from app.models.case import Case
from app.models.document import Document, DocumentType
import hashlib
import json


def _replace_in_paragraph(paragraph, mapping: Dict[str, str]) -> None:
    """
    Replace placeholders while keeping template formatting as much as possible.
    - Prefer run-level replace (keeps font/style).
    - Fallback: if placeholder spans runs, do a safe paragraph-level rebuild
      (may affect mixed formatting in that paragraph).
    """
    # 1) run-level replace (preserve formatting)
    hit = False
    for run in paragraph.runs:
        if not run.text:
            continue
        new_text = run.text
        for k, v in mapping.items():
            if k in new_text:
                new_text = new_text.replace(k, v)
        if new_text != run.text:
            run.text = new_text
            hit = True

    # 2) fallback if placeholders span multiple runs
    # (e.g. "{{seller_name}}" bị tách thành nhiều run)
    if not hit:
        full = paragraph.text
        new_full = full
        for k, v in mapping.items():
            if k in new_full:
                new_full = new_full.replace(k, v)

        if new_full != full:
            # rebuild paragraph text: keep first run formatting if exists
            if paragraph.runs:
                paragraph.runs[0].text = new_full
                for r in paragraph.runs[1:]:
                    r.text = ""
            else:
                paragraph.add_run(new_full)


def _replace_everywhere(doc, mapping: Dict[str, str]) -> None:
    for para in doc.paragraphs:
        _replace_in_paragraph(para, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, mapping)


def generate_contract(case: Case, db: Session) -> Document:
    doc_type = DocumentType.CONTRACT_TRANSFER
    new_hash = compute_content_hash(case, doc_type)

    latest = db.execute(
        select(Document)
        .where(Document.case_id == case.id, Document.doc_type == doc_type)
        .order_by(Document.version.desc())
        .limit(1)
    ).scalar_one_or_none()

    if latest and latest.content_hash == new_hash and latest.file_path and Path(latest.file_path).exists():
        return latest

    new_version = (latest.version + 1) if latest else 1

    doc = DocxDocument("templates/contract_transfer.docx")

    prop = case.property
    data = {
        "{{transfer_price}}": f"{case.transfer_price:,}".replace(",", ".") if case.transfer_price else "",
        "{{signing_date}}": str(case.signing_date) if case.signing_date else "",

        "{{property_address}}": prop.address if prop else "",
        "{{property_map_sheet_no}}": prop.map_sheet_no if prop else "",
        "{{property_parcel_no}}": prop.parcel_no if prop else "",
        "{{property_area_m2}}": str(prop.area_m2) if (prop and prop.area_m2 is not None) else "",
        "{{property_certificate_no}}": prop.certificate_no if prop else "",

        "{{seller_name}}": "",
        "{{seller_cccd}}": "",
        "{{seller_address}}": "",
        "{{buyer_name}}": "",
        "{{buyer_cccd}}": "",
        "{{buyer_address}}": "",
    }

    for cp in (case.parties or []):
        role = getattr(cp.role, "name", None)
        if role == "SELLER":
            data["{{seller_name}}"] = cp.party.full_name or ""
            data["{{seller_cccd}}"] = cp.party.cccd or ""
            data["{{seller_address}}"] = cp.party.address or ""
        elif role == "BUYER":
            data["{{buyer_name}}"] = cp.party.full_name or ""
            data["{{buyer_cccd}}"] = cp.party.cccd or ""
            data["{{buyer_address}}"] = cp.party.address or ""

    _replace_everywhere(doc, data)

    out_dir = Path("data/files") / case.code
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"contract_v{new_version}.docx"
    doc.save(str(out_path))

    record = Document(
        case_id=case.id,
        doc_type=doc_type,
        version=new_version,
        file_path=str(out_path),
        content_hash=new_hash,
    )
    db.add(record)
    db.flush()
    db.refresh(record)
    return record



def _canonical_case_payload(case: Case, doc_type: DocumentType) -> dict:
    prop = case.property
    parties = []
    for cp in (case.parties or []):
        parties.append({
            "role": getattr(cp.role, "name", str(cp.role)),
            "cccd": cp.party.cccd,
            "full_name": cp.party.full_name,
            "address": cp.party.address,
            "phone": cp.party.phone,
            "cccd_issue_date": str(getattr(cp.party, "cccd_issue_date", "") or ""),
            "cccd_issue_place": getattr(cp.party, "cccd_issue_place", "") or "",
        })

    # sort để hash ổn định (không phụ thuộc thứ tự load)
    parties.sort(key=lambda x: (x["role"], x["cccd"] or ""))

    return {
        "doc_type": doc_type.value if hasattr(doc_type, "value") else str(doc_type),
        "case": {
            "id": case.id,
            "code": case.code,
            "case_type": case.case_type.value if hasattr(case.case_type, "value") else str(case.case_type),
            "signing_date": str(case.signing_date or ""),
            "transfer_price": str(case.transfer_price or ""),
        },
        "property": None if not prop else {
            "address": prop.address,
            "map_sheet_no": prop.map_sheet_no,
            "parcel_no": prop.parcel_no,
            "area_m2": str(prop.area_m2 or ""),
            "certificate_no": prop.certificate_no,
        },
        "parties": parties,
    }


def compute_content_hash(case: Case, doc_type: DocumentType) -> str:
    payload = _canonical_case_payload(case, doc_type)
    raw = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
    return hashlib.sha256(raw).hexdigest()